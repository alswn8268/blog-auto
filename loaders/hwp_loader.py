"""문서 로더 — HWP/HWPX 1순위, 실패 시 순수 파이썬 파서로 폴백 (명세 2.1).

파서 라이브러리는 환경에 따라 설치 여부가 갈리므로 모두 지연 임포트한다.
`load_document()`가 확장자를 보고 알맞은 로더를 고르는 단일 진입점이다.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".hwp", ".hwpx", ".pdf", ".docx", ".txt", ".md"}


@dataclass
class LoadedDocument:
    """파서 종류와 무관하게 파이프라인이 받는 공통 형태."""

    text: str
    source_file: str
    tables: list[str] = field(default_factory=list)
    loader: str = "unknown"


class DocumentLoadError(RuntimeError):
    """어떤 파서로도 문서를 읽지 못했을 때."""


# ── HWP / HWPX ─────────────────────────────────────────────────────────────


def load_hwp(filepath: str | Path) -> LoadedDocument:
    """1순위: rhwp의 LangChain 연동 로더. HWP/HWPX 모두 지원하고 빠르다.

    주의: rhwp는 비교적 최근 공개된 패키지라 API가 바뀔 수 있다.
    실제 사용 전 PyPI/GitHub의 최신 사용법을 한 번 더 확인할 것.
    """
    from rhwp.integrations.langchain import HwpLoader  # 지연 임포트

    docs = HwpLoader(str(filepath)).load()  # LangChain Document 리스트
    text = "\n\n".join(d.page_content for d in docs)
    return LoadedDocument(text=text, source_file=str(filepath), loader="rhwp")


def load_hwp_fallback(filepath: str | Path) -> LoadedDocument:
    """대안: hwp-hwpx-parser (순수 Python, JVM·Windows 불필요).

    표를 마크다운으로 뽑아주므로 별표·별지서식이 많은 문서에서 결과가 더 깨끗할 수 있다.
    """
    from hwp_hwpx_parser import Reader  # 지연 임포트

    with Reader(str(filepath)) as r:
        return LoadedDocument(
            text=r.extract_text(),
            tables=list(r.get_tables_as_markdown() or []),
            source_file=str(filepath),
            loader="hwp-hwpx-parser",
        )


# ── 그 밖의 포맷 ───────────────────────────────────────────────────────────


def load_pdf(filepath: str | Path) -> LoadedDocument:
    from pypdf import PdfReader  # 지연 임포트

    reader = PdfReader(str(filepath))
    text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
    return LoadedDocument(text=text, source_file=str(filepath), loader="pypdf")


def load_docx(filepath: str | Path) -> LoadedDocument:
    import docx  # 지연 임포트 (python-docx)

    document = docx.Document(str(filepath))
    text = "\n".join(p.text for p in document.paragraphs)
    return LoadedDocument(text=text, source_file=str(filepath), loader="python-docx")


def load_text(filepath: str | Path) -> LoadedDocument:
    path = Path(filepath)
    return LoadedDocument(
        text=path.read_text(encoding="utf-8", errors="replace"),
        source_file=str(path),
        loader="plaintext",
    )


# ── 단일 진입점 ────────────────────────────────────────────────────────────

HWP_LOADERS = (load_hwp, load_hwp_fallback)


def load_document(filepath: str | Path) -> LoadedDocument:
    """확장자에 맞는 로더로 문서를 읽는다.

    HWP/HWPX는 1순위 로더가 실패하면 폴백 파서를 순서대로 시도한다.
    1주차에 실제 샘플로 두 파서를 비교해보고, 결과가 더 깨끗한 쪽을
    HWP_LOADERS 앞자리에 두면 된다.
    """
    path = Path(filepath)
    if not path.exists():
        raise DocumentLoadError(f"파일이 없습니다: {path}")

    suffix = path.suffix.lower()
    if suffix in {".hwp", ".hwpx"}:
        errors: list[str] = []
        for loader in HWP_LOADERS:
            try:
                return loader(path)
            except Exception as exc:  # 파서마다 예외 타입이 제각각이라 광범위하게 잡는다
                logger.warning("%s 실패 (%s): %s", loader.__name__, path.name, exc)
                errors.append(f"{loader.__name__}: {exc}")
        raise DocumentLoadError(f"HWP 파싱 실패 ({path.name})\n" + "\n".join(errors))

    if suffix == ".pdf":
        return load_pdf(path)
    if suffix == ".docx":
        return load_docx(path)
    if suffix in {".txt", ".md"}:
        return load_text(path)

    raise DocumentLoadError(f"지원하지 않는 확장자입니다: {suffix}")


def iter_documents(directory: str | Path):
    """디렉터리 안의 지원 문서를 파일명 순으로 순회한다."""
    for path in sorted(Path(directory).rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES:
            yield path
